"""Module containing the core Class objects and functions required to support execution of the main position description program. The items contained in this module include the separate functions for each of the various tasks associated with the generation of a standard formatted Microsoft Word document with static png image of an interactive plotly radar chart that also pushes the document to a Google Drive account and folder labeled 'Position Descriptions' and pushes the radar chart to a Plotly Chart Studio account, generating a unique access url.

Created by: Babila R. Lima
Creation date: September 9, 2019
"""


from datetime import datetime
import os
import time

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import chart_studio.plotly as py
import plotly.graph_objects as go
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
from tqdm import tqdm
import yaml



###########################################################################################
                                    ## Position Class ##
###########################################################################################

organization_info = """\tAs an agency the Department of General Services manages the vertical and mobile assets for the City of Baltimore, providing the critical infrastructure for the operation of government agencies, including those providing direct public services to citizens.   Organizational units are organized around these core responsibilities. 
\tThe Business Process Improvement Office (BPIO) is a unit of the Department of General Services. We are a performance management consultancy to organizational units and services within the agency.  BPIO drives continuous process improvement for service delivery helping across a range of domains and tasks including but not limited to: research, policy development, key performance indicator development, project implementation, statistical analysis and task automation. We help create and communicate quantifiable advancements in government efficiency, improving operating time, capital and human resource allocations or requirements.\tBPIO has three core focus areas: 1) completely reducing the agency’s dependency on paper, making all processes paperless by 2024, 2) furthering the BPIO Extension Project via partnerships with relevant private industry, public sector and academic entities to promote low cost process, information technology or project-based policy solutions or experiments and information sharing to create a sustainable highly skilled pipeline to public service that exponentially increases problem solving capacity, and 3) automating, to the maximum extent possible, all repetitive analysis, communications and administrative tasks throughout the agency and establishing Python as the lingua franca for municipal government process improvement. 
"""

class Position:
    """The object contains the core common elements of a work position to be advertised and recruited for.
    
    Attributes
    ----------
    title (str): The name of the position or job title
    
    about_job (str): The descriptive or background text about the position.
    
    about_org (str): The description and background about the agency or office.
    
    workplan (dict): Dictionary with key, value pairs for the job task or compentency area and 
            the percentage of time anticipated or necessary to be spent in that area to successfully  
            execute the job.
    
    """
    def __init__(self, title, about_job, workplan):
        """The constructor for the job position class.
        
        Attributes
        ----------
        title (str): The name of the position or job title

        about_job (str): The descriptive or background text about the position.

        about_org (str): The description and background about the agency or office.

        workplan (dict): Dictionary with key, value pairs for the job task or compentency area and 
                the percentage of time anticipated or necessary to be spent in that area to successfully  
                execute the job.
        
        division (str): The name of the work unit or division in which the position is located.
        """
        self.title = title
        self.about_job = about_job
        self.workplan = workplan
        self.division = 'BPIO: Business Process Improvement Office'
        self.about_org = organization_info

        
        
###########################################################################################
                            ## Methods for Accessing Position Class  ##
###########################################################################################

def get_position_data(file):
    """Parse yaml document containing data on each available position.
    Function parses a yaml document into stream and laods into a Python object.
    
    Parameters
    ----------
    file: Str
        Yaml file containing the information on each position. Each position in 
        the Yaml file has an attribute corresponding to the attributes for the 
        Position class. 
        
    Returns
    -------
    dict:  Returns dictionary of the parsed Yaml file. The dictionary has key value 
           pairs for the position as keys and the corresponding attirbutes for the 
           position as values. 
           
    Example
    -------
    >>> get_position_data(file='data/position_data.yaml') # returns dict of positions & attributes
    """
    stream = open(file, mode='r')
    data = yaml.safe_load(stream)
    
    return data
    

def generate_jobs(data):
    """Generate new instance of the Position class for each job found in the data.
    
    Function uses the parsed data from yaml file containing positions and position data
    and for each job it creates an instance of the Position class, storing each 
    instances in a dictionary.
    
    Parameters
    ----------
    data:   dict
         dictionary of the parsed Yaml file. The dictionary has key value pairs for 
         the position as keys and the corresponding attirbutes for the position as values.
         
    Returns
    -------
    dict:  Returns dictionary with key value pairs for the job and attributes of the job.
    
    Example
    -------
    >>> generate_jobs(data=position_data)
    
    >>> generate_jobs(get_position_data(file='data/position_data.yaml'))
    """
   
    jobs = {}
    for job in data:
        jobs[job] = Position(title=data[job]['title'],
                            about_job=data[job]['description'],
                            workplan=data[job]['workplan'])
    
    return jobs

def show_jobs(file):
    """Return list of all positions available.
    
    Parameters
    ----------
    file:  Str
        Yaml file containing the information on each position. Each position in 
        the Yaml file has an attribute corresponding to the attributes for the 
        Position class. 
    
    Returns
    -------
    list:  Returns list of positions available in Yaml file containing positions. 
    
    Example
    -------
    >>> show_all_positions(file='data/positions_data.yaml')
        
    """
    
    try:
        data = get_position_data(file)
    
    except Exception as e:
        return e
    
    try:
        positions = []
        for position in data.keys():
            positions.append(position)
    
        return positions
    
    except Exception as e:
        return e
    

def show_position_titles(file):
    """Return list of all position job titles.
    
    Parameters
    ----------
    file:  Str
        Yaml file containing the information on each position. Each position in 
        the Yaml file has an attribute corresponding to the attributes for the 
        Position class. 
    
    Returns
    -------
    list:  Returns list of job titles for each position available in the Yaml file 
    containing positions. 
    
    Example
    -------
    >>> show_position_titles(file='data/positions_data.yaml')
        
    """
    
    try:
        data = get_position_data(file)
        
    except Exception as e:
        return e
    
    try:
        jobs = generate_jobs(data)
    
    except Exception as e:
        return e
    
    try:
        titles = []
        for job in jobs:
            titles.append(jobs[job].title)
            
        return titles
    
    except Exception as e:
        return e
    
    
    

###########################################################################################
                                    ## Helper Functions ##
###########################################################################################


def read_position_file(filename):
    """Parse yaml document containing data on a single position.
    Function parses a yaml document into stream and laods into a Python object.
    
    Parameters
    ----------
    filename: Str
        Yaml file containing the information on each position. Each position in 
        the Yaml file has an attribute corresponding to the attributes for the 
        Position class. 
        
    Returns
    -------
    dict:  Returns dictionary of the parsed Yaml file. The dictionary has key value 
           pairs for the position as keys and the corresponding attirbutes for the 
           position as values. 
           
    Example
    -------
    >>> read_position_file(filename='data/some_position_title.yaml') # returns dict of positions & attributes
    """
    stream = open(filename, mode='r')
    data = yaml.safe_load(stream)
    
    return data



def create_position_object(data):
    """Create easy access dictionary with custom keys and with values set to elements of the position file.
    
    Parameters
    ----------
    data:    dict
        Dictionary object returned from the read_position_file() method which reads the yaml file where position 
        data is stored. 
    
    Returns 
    -------
    dict:    dict
        Returns a dictionary with the following keys: ['title','job_summary','job_expectations',
        'job_scope', 'job_activities', 'breakdown'].
    
    Example
    -------
    >>> create_position_object(data = position_data)
    >>> create_position_object( read_position_file('data/some_position_title.yaml') ) # takes function call as an argument
    """
    position = {} 
    for element in data.values():
        position['title'] = element['title']
        position['job_summary'] = element['summary']
        position['job_expectations'] = element['expectations']
        position['job_scope'] = element['scope']
        position['job_activities'] = element['activities']
        position['breakdown'] = element['workplan']
        
    return position


def get_position(yaml_file):
    """Reads and converts data from yaml file into a dictionary with keys exposing core elements of a position description.
    
    Parameters
    ----------
    yaml_file:   Str
        Name of the yaml file containing the data on a target position. The yaml file has sections
        corresponding to core elements of a position description. 
    
    Returns
    -------
    dict:    dict
        Returns a dictionary with the following keys: 
        ['title','job_summary','job_expectations','job_scope', 'job_activities', 'breakdown'].
    
    
    Examples
    --------
    >>> get_position(yaml_file='data/some_position_name.yaml')
    """
    try:
        data = read_position_file(yaml_file)
        
        if data:
            try:
                job_details = create_position_object(data=data)
                
                return job_details
            
            except:
                print('Error: create_position_object() method failed')
    
    except:
        print(f'Error: failed attempting to handle and read {yaml_file}')
        
        
        
# generate chart
class Error(Exception):
    """Base class for throwing exceptions"""
    pass

class ValuesDontSumTo100(Error):
    """Workplan values fail to sum to 100."""
    pass

class ValuesNotTypeInt(Error):
    """Workplan values not all type int."""
    pass

class InvalidParameterType(Error):
    """Value passed to workplan parameter not type dict."""
    pass

def validate_workplan(workplan):
    """Checks whether the workplan argument is a dictionary and whether or not percentages associated 
    with job tasks total to 100. 
    
    Parameters
    ----------
    workplan:  Dict
            Dictionary with key, value pairs for the job task or compentency area and 
            the percentage of time anticipated or necessary to be spent in tha area for the 
            job.  Example: {'memo writing':20, 'answering phones':20, 'drinking coffee':60}
    
    Returns
    -------
    Boolean:  Returns True if sum of values is 100, Raises error if not. 
    """
    
    if isinstance(workplan, dict):
        
        if all(type(value) is int for value in workplan.values()):
            
            if sum(workplan.values()) == 100:
                return True
   
            else:
                error_msg = 'The workplan parameter values should total 100, instead sums to {}.'.format(sum(workplan.values()))
                raise ValuesDontSumTo100(error_msg)
        else:
            
            error_msg = ('All workplan values must be integers, instead received type(s): {}.'
                         .format([type(value) for value in workplan.values() if not isinstance(value, int)]))
            raise ValuesNotTypeInt(error_msg)
    
    else:
        error_msg = 'The workplan parameter expects a dictionary, but received {}.'.format(type(workplan))
        raise InvalidParameterType(error_msg)
        


        
## TODO: style hovertext to show workarea: BPM <br> pct: 23%  formatting

def create_radarchart(position):
    """Creates radar chart for a position to visualize the percentage of time 
    and effort required to be spent in each competency or task area for a job.
    
    Note:
    The function relies on a validate_workplan() method to ensure that the sum of 
    values associated with elements of a workplan equal to 100. 
    
    Parameters
    ----------
    position: dict
        A dictionary with  keys for the core elements of a position description like: 
        ['title','job_summary','job_expectations','job_scope', 'job_activities', 'breakdown']
             
        
    Returns
    -------
    Dict:   
        Returns object of type: plotly.graph_objs._figure.Figure which is a Plotly figure dictionary used 
        for creating interactive plots that has keys: 'data', 'layout' and 'frames'.

    
    Example
    -------
    >>> create_radarchart(position=SecretaryII)
    
    >>> create_radarchart(GrantManager)
    """
    
    jobtitle, workplan = position['title'], position['breakdown'] 
    
    if validate_workplan(workplan):
        
        try:
            
            responsibilities = list(workplan.keys())
            r_values = list(workplan.values())
            axis_range = [0,(max(r_values) + 5)]


            data = [go.Scatterpolar(
            r = r_values,
            theta = responsibilities,
            fill = 'toself',
            fillcolor = '#a1d99b',
            line = {'color':'#a1d99b'},
            #hoverinfo = 'r+theta'
            )]

            layout = {
                'title':'{}<br>Responsibilities'.format(jobtitle),
                'font':{'color':'#a1d99b'},
                'polar':{
                    'radialaxis':{'visible':True,
                                 'color':'black',
                                 'linecolor':'green',
                                 'range': axis_range} 
                },
                'showlegend':False,
                'hovermode':'closest'
            }
            
            ## Figure object has write method properties useful for generating image files
            fig = go.Figure(data=data, layout=layout)
            
            return fig
        
        except Exception as e:
            return e
    
    else:
        error_msg = 'The workplan parameter expects a dictionary, but received {}.'.format(type(workplan))
        raise InvalidParameterType(error_msg)
        
        
        
def ensure_filepath_exists(filename):
    """Checks for, awaits and  verifies existence of a file path to register in operating 
    system after some preceding file generating function call. 
    
    Parameters
    ----------
    file:   Str
        The name of target file recently generated by a separate, outside function. 
    
    Returns
    -------
    Str:
        Filepath to file. The path returned is one verified to exist on the operating system and 
        accessible for manipulation or access immediately.  
        
    Examples
    --------
    >>> ensure_filepath_exists(file='data/charts/some_image_file.png') 
    
    >>> higher_level_function(param1 = some_argument,
                              input_file = ensure_filepath_exists('data/charts/img.png'),
                              param3 = some_other_argument
    )
    """
    
    if os.path.exists(filename):
        
        path = filename
    
    else:
        
        try:
            wait_seconds_iterable = [1,2,3,4,5]
            for second in wait_seconds_iterable:
                time.sleep(second)
                
                if os.path.exists(filename):
                    
                    path = filename
                    
                    break
                        
        except Exception as e:
            raise e
            
    
    if not os.path.exists(filename):
        print(f'Error: failed to ensure existence of filepath for {filename}')

        
    return path



def save_plot(position, extension='png'):
    """Saves plotly radarchart as an image file returning a tupple with both the filename and chart figure object. 
    
    Parameters
    ----------
    position:   dict
        Returns a dictionary with the core elements associated with a position as keys like: 
        ['title','job_summary','job_expectations','job_scope', 'job_activities', 'breakdown'].
    
    
    extension:  Str, optional
        File extension for type of image file to which the plot should be saved. (The
        default is 'png' which will save the chart as a png file.)
        
        
    Returns
    -------
    Str:
        The name of the png file where the chart is saved. 
        
    Dict:   
        
        Returns object of type: plotly.graph_objs._figure.Figure which is a Plotly figure dictionary 
        used for creating interactive plots that has keys: 'data', 'layout' and 'frames'.
    
    Examples
    --------
    >>> save_plot(position=MasterGardener) # returns a png file with current date in filename and plot figure object.
    >>> save_plot(DataArchitect, 'png')
    """
    timestamp = datetime.now().strftime('%Y-%m-%d')
    
    folder = 'data/charts'
    file = f"{position['title']} {timestamp}.{extension}"
    filename = os.path.join(folder,file)
    
    if not os.path.exists(folder):
        os.mkdir(folder)
        
    else:
        pass
    
    
    try:
        figure = create_radarchart(position=position)
        
        if figure:
            
            try:
                fname = filename
                figure.write_image(file=fname, format=extension)
                
                if not os.path.exists(fname):
                    ensure_filepath_exists(fname)
                else:
                    pass
                
                
                return fname, figure
            
            except Exception as e:
                raise e
                print(f'Error: the plotly Figure write_image() method or ensure_filepath_exists() method failed with {fielname}' )
            
    except Exception as e:
        raise e
        print('Error: create_radarchart() method failed to generate ploty chart')
        
        
        
def strip_filename(image_filename, directory='data/charts/'):
    """Strips directory and extension from filename. 
    
    Shaving the filename down allows for sending only parts of the filename relevant to 
    what should appear associated with the chart in the plotly account.
    
    Parameters
    ----------
    image_filename:   Str
        Name of the image file for the saved chart illustrating the workplan breakdown.
        
    directory:        Str, optional
        Name of the folder where the chart image files are stored. 
        
    Returns
    -------
    Str:
        File name with directory and extension striped. For example 'data/charts/Analyst1.png' becomes
        Analyst1
    
    Examples
    --------
    >>> strip_filename('data/charts/some_position_name 2019-09-10.png') 
    
    """
    
    name = image_filename.split('.')[0].split(directory)[1]
    return name



def publish_chart(filename, figure, fileoption='overwrite'):
    """Creates a unique url for chart in plotly and publishes plot to chart studio account. 
    
    The function pushes a chart to the account and sets the sharing parameter to 'public' so that is publicly viewable. It also publishes the chart with default setting to overwrite a file if 
    there is a file name clash. 
    
    Parameters
    ----------
    filename:    Str
        The name of the png file where the chart is saved. 
    
    figure:      dict
        Object of type: plotly.graph_objs._figure.Figure which is a Plotly figure dictionary 
        used for creating interactive plots that has keys: 'data', 'layout' and 'frames'.
        
    fileoption:  Str, optional
        The option for how Plotly should handle future calls to the function for the same chart. Valid
        values include: 'new', 'overwrite', 'extend', and 'append'. They have the following actions: 
        ('new' | 'overwrite' | 'extend' | 'append') -- 'new' creates a
        'new': create a new, unique url for this plot
        'overwrite': overwrite the file associated with `filename` with this
        'extend': add additional numbers (data) to existing traces
        'append': add additional traces to existing data lists 
    
    Returns
    -------
    Str:   
        Returns the unique url for the chart as a string. Example: 'https://plot.ly/~brl1906/1540/'
        
        
    Examples
    --------
    >>> publish_chart(filename='data/charts/automMechanic 2019-05-11.png', figure=fig)
    >>> publish_chart(fname, fgr)
    >>> publish_chart(fname, fig, fileoption='new')
    """
    try:
        stripped = strip_filename(filename)
        
    except:
        print(f'Error: function failed to shorten filename {filename}')
        
    if stripped:
        
        try:
            plot_options = {'sharing':'public', 'filename':stripped}
            url =  py.plot(figure_or_data=figure, auto_open=False, fileoption=fileoption, plot_options=plot_options)
            
            return url
                    
        except Exception as e:
            
            raise e
            print('Error: failed to create unique url for chart and push to plotly account')
            
            

def generate_plot(position):
    """Handles workplan related data associated with a position and creates, saves and publishes to plotly chart studio account a radar chart illustrating a breakdown of the workplan.
    
    The function uses the following default settings and outputs: 
    1) It saves the chart image as a png filetype
    2) It creates or puts the image file in a directory named 'data/charts'
    3) The published chart sets the name of the plot in the plotly studio account to the name of the image file stripped of the directory and file extension.
    4) The published chart sharing setting is 'public' allowing anyone with the link to access the interactive plot.
    
    Parameters
    ----------
    position: dict
        A dictionary with data on the core elements of a position description that includes core elements 
        as dictionary keys like: ['title','job_summary','job_expectations','job_scope', 'job_activities', 'breakdown'].
    
    
    Returns
    -------
    Str:
        The name of the png file where the chart is saved.
    
    Str:   
        Returns the unique url for the published chart as a string. Example: 'https://plot.ly/~brl1906/1540/'
    
    
    Examples
    --------
    >>> generate_plot(executive_assistant)
    >>> generate_plot(position=get_position('data/marketing intern.yaml'))
    """
    
    ## generate & save plot as png file
    try:
        fname,fig = save_plot(position) # saved plot function returns tupple
    
    except Exception as e:
        
        raise e
        print("Error: failed to generate and save plot for {}".format(position['title']))
     
    ## publish the plot
    if fname and fig:
        try:
            url = publish_chart(filename=fname, figure=fig)
            
            return fname, url
        
        except Exception as e:
            
            raise e
            print("Error: failed to publish chart to plotly account for position {}".format(position['title']))
            
            

              
def insert_image(document, image, image_width=3.6, image_height=3):
    """Inserts an image into a word document with default 3x3 size settings.
    
    Parameters
    ----------
    document:     docx.document.Document
        An object containing an initialized word document object. The document object is not intented to 
        be constructed directly but is initialized as an object to manipulate, add to and futher
        construct a Microsoft Word document with other elements as needed. A Document object loaded from docx, 
        where docx can be either a path to a .docx file (a string) or a file-like object. More on the usage of 
        python-docx Document can be found here: https://python-docx.readthedocs.io/en/latest/api/document.html?highlight=document
    
    
    image:        Str
        The filepath to the location of the radar chart. 
    
    image_width:  int, optional
        Setting for the width of the image in inches on a Microsoft Word document. (Default set to 3.6 inches)
        
    image_height:  int, optional
        Setting for the height of the image in inches on a Microsoft Word document. (Default set to 3 inches)
    
    Returns
    -------
    int: 
        Return code 0 for success and 1 for failure.
    
    Examaples
    ---------
    >>> insert_image(document=docx_object, picture=imgage_file)
    >>> insert_image(docx_object, img_file, 2.5, 2.5) # shrinks image size on document to 2.5 by 2.5.
    """
    try:
        document.add_picture(image, width=Inches(image_width), height=Inches(image_height))
        return 0 # success
    
    except Exception as e:
        raise e
        print(f'Error: insertion of {picture} image failed.')
        return 1 # error 

    
    
def create_document(position, img_file, about_org=organization_info):
    """Converts data on the position into a standard formatted Microsoft Word document of a position 
    description. 
    
    The function accesses data from relevant keys in the position argument to create a word document
    with sections including: title, 'About Organization', 'Position Summary', 'Expectations & Outcomes',
    'Position Scope', 'Activities & Deliverables'. It also includes an image of a radar chart illustrating
    position workplan and time allocation across task areas as the last item on the document. 
    
    
    Parameters
    ----------
    position:  dict
        A dictionary with  keys for the core elements of a position description like: 
        ['title','job_summary','job_expectations','job_scope', 'job_activities', 'breakdown']
    
    img_file:  Str,
        The file location for the image of the chart illustrating the workplan breakdown.
    
    about_org: Str, optional
        Standard descriptive introdcution about the agency. (Default is pulled from helper module, Position class)
    
    Returns
    -------
    Str:
        Filepath for newly created Microsoft Word document for the position description. In generating a
        filename for the document the function uses the position title and a timestamp from the time the 
        functin is run to generate the document.   For example, a position with the title of Grants Manager 
        run Christmas eve would generate a document file name of 'Grants Manager 12-25-2019.docx'
    
    Examples
    --------
    >>> create_document(marketing_specialist, 'data/charts/markSpec 2019-15-23.png')
    
    >>> create_document(position = marketing_specialist,
                        img_file = 'data/charts/markSpec 2019-08-22.png',
                        about_org = 'Text about organization as string or pass variable.'
                        )
    """
    timestamp = datetime.now().strftime('%Y-%m-%d')
    filename = f"{position['title']} {timestamp}.docx"
    
    doc = Document()
    ## Title Section of Document ##
    title = doc.add_paragraph(text="Position Title: {}\n".format(position['title']))
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ## About Organization Section ##
    org_section = doc.add_paragraph(text='')
    org_section.add_run('About Organization\n').bold = True
    org_section.add_run(text=about_org, style=None)

    ## Position Summary Section ##
    position_section = doc.add_paragraph(text='')
    position_section.add_run(text='Position Summary\n').bold = True
    position_section.add_run(text='General Purpose:\n').italic = True
    position_section.add_run(text=position['job_summary'], style=None)

    ## Professional Expectations & Outcomes Section ##
    expectations_section = doc.add_paragraph(text='')
    expectations_section.add_run(text='Professional Outcomes & Expectations:\n').italic = True
    expectations_section.add_run(text=position['job_expectations'], style=None)

    ## Position Scope Section ##
    scope_section = doc.add_paragraph(text='')
    scope_section.add_run(text='Position Scope:\n').italic = True
    scope_section.add_run(text=position['job_scope'], style=None)

    ## Activities & Deliverables Section ##
    activities_section = doc.add_paragraph(text='')
    activities_section.add_run(text='Activities & Deliverables\n').bold = True
    activities_section.add_run(text=position['job_activities'] ,style=None)
                              
    
    if os.path.exists(img_file):
        
        try:
            
            insert_image(document=doc, image=img_file)
            
            image_block = doc.paragraphs[-1]
            image_block.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        except Exception as e:
            raise e
            print('Error: failed to insert into document image file: {}'.format(img_file))
            
    else:
        print('No image to attach to document. {} does not exist'.format(img_file))
    
    
    try:
        
        doc.save(filename)
        
        ## wait for file to save and register in file system before exiting
        ## functions requesting to do something with the returned document fail if called
        ## immediately upon completion of the save if file hasn't yet registered.
        
        if os.path.exists(filename):
            
            path = filename

        else:
            
            wait_seconds_iterable = [1,2,3,4,5]
            for second in wait_seconds_iterable:
                print(second)
                time.sleep(second)

                if os.path.exists(filename):
                    path = filename
                    break

        return path    
    
    except:
        print(f'Error: file {filename} failed to save.')
        
    

def google_authentication():
    """Logs into Google account and returns a connected and authorized object acting as the drive.
    
    Function login happens via the pydrive module's wrapper class for the oauth2client library in 
    the google-api-python-client. It uses settings and data from a credentials file for authentication.
    
    Note:
    The default settings for pydrive allow authentication but prompt the user with a browser forcing them to 
    select which Google account and to click approve each time the function is run.  However, in order to 
    automate authentication and login with Google such that it can be done through a remotely hosted script 
    or command line application, the program stores authentication credentials in a separate yaml settings file. 
    
    After authenticating for the first time via the browser, a credentials.json file is automatically generated
    with a refresh token.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    pydrive.drive.GoogleDrive:
        Represents object of a sucessful authentication to user's Google account. The googledrive
        object is of type pydrive.drive.GoogleDrive
    
    """
    try:
        login = GoogleAuth()
        login.LocalWebserverAuth()
        drive = GoogleDrive(login)
        
        return drive
    
    except:
        print('Login failed')
        
        return 1
    
    
## :: WARNING ::
## Running this function may lead to duplicate file folders being generated in Google Drive account ##
def create_unique_google_folder(drive, name):
    """Creates a unique Google Drive folder and returns the folder id. 
    
    A Google Drive folder is essentially just another Google Drive file. Using the
    id of a folder, files can be added to specific folders by linking the upload of 
    a new file to a parent folder by it's ID.  
    
    WARNING: If this function is run more than once, it will create multiple folders
    with the same folder name in Google Drive--same 'title' property--but each with
    different and unique folder ID. This will lead to significant confusion, rendering
    what looks like multiple of the same folder that are actually different folders. 
    
    Parameters
    ----------
    drive: Str
        The object is a sucessful authentication to user's Google account. The googledrive
        object is of type pydrive.drive.GoogleDrive
    
    
    name:  Str
        Name of folder to be added to Google Drive
        
    Returns
    -------
    Str:
        Unique ID for the folder in Google Drive. The ID allows for targeting a specific
        folder for an action like, deletion, renaming or adding a specific file to a 
        target folder. 
        
    Examples
    --------
    >>> create_unique_google_folder(name='Position Descriptions') 
    
    """
    folder_metadata = {
        'title': name,
        'mimeType': 'application/vnd.google-apps.folder'
    }
    folder = drive.CreateFile(folder_metadata)
    
    try:
        folder.Upload()
        
    except Exception as e:
        
        raise e
        print(f'Error: Upload() method failed to create google drive folder for {name}')

    folderID = folder['id']
    
    return folderID    


def get_googledrive_folderID(division):
    """Access configuration file to get the ID for a target Google Drive folder.
    
    Parameters
    ----------
    division:    Str
        Identifies which division or unit of the organization and relevant folder is
        being targeted in the configuration file. While not returned from the function, 
        having access to this element allows for future dynamic generation of file or 
        folder names if there is broad use across divisions. (Default is 'bpio') 
        
        Options include:
            bpio
            fleet
            facilities
            capitalProjects
            administration
        
    Returns
    -------
    Str:
        Google Drive folder ID. Having the ID allows for specific actions on a target folder 
        in Google Drive such as deletion, renaming, or adding files to a target folder.
        
    Example
    -------
    >>> get_googledrive_folderID('data/google_folders.yaml')
    
    """
    stream = open('data/google_folders.yaml')
    data = yaml.safe_load(stream)
    folderID = data[division]['folderID']
    
    return folderID



def googledrive_upload(googledrive, file, folder):
    """Converts a document to a Google Drive file and sends it to user's Google Docs account.
    
    Note: The SetContentFile() method opens and reads the target file file. If the metadata argument
    is not specified, the method automatically sets a default for 'title' and 'mimeType' values. 
    Calling the Upload() method both uploads and closes the file.  Additional documentation on pydrive 
    SetContentFile() and Upload() methods can be found here: https://pythonhosted.org/PyDrive/index.html
    
    Parameters
    ----------
    googledrive:  Str
        String represents object of a sucessful authentication to user's Google account. The googledrive
        object is of type pydrive.drive.GoogleDrive
        
    file:         Str
        Filepath for the document to be opened, read and used to create a Google Doc.
    
    folder:       Str
        Unique id for a Google Drive direcotry that allows for targeting a specific folder for creation, 
        deletion, or for adding files to a target directory. Example folder id string could look like
        the following: 169W-OsOEVpJu2vSmhhbQ2a-ZqMVCgegG as an example
    
    
    Returns
    -------
    int:
        Returns return code 0 if sucessful and 1 if failed.
    
    Examples
    --------
    >>> googledrive_upload(goooledrive='Gauth', file='MyPositionDescriptionFile.docx')
    
    """
    if folder == None:
    
        try:
            document_metadata = {'title': f'{file}'}
            newfile = googledrive.CreateFile(document_metadata)
            newfile.SetContentFile(file)

        except Exception as e:
            raise e
            print(f"Error: google drive SetContentFile() method failed to generate a new google document object from  {file}" )
            return 1

        try:
            newfile.Upload() # upload file to google drive account 
        except Exception as e:
            raise e
            print(f"Error: google drive Upload method failed to upload {file}")
            return 1

        print(f"File uploaded successfully: {file}")
        return 0
    
    else:
        
        folder_metadata = {'title':'Position Descriptions','mimeType': 'application/vnd.google-apps.folder'}
        folderID = folder
        
        try:
            ## create target folder in google drive as parent object & upload new file to the parent 'file' folder
            document_metadata = folder_metadata
            newfile = googledrive.CreateFile({'parents': [{"kind": "drive#fileLink", "id": folderID}]})
            
            newfile.SetContentFile(file)
        
        except Exception as e:
            raise e
            print('Error: google drive SetContentFile() method failed to create new google drive folder named {}'.
                  format(folder_metadata['title']))
        
        try:
            newfile.Upload()
        except Exception as e:
            raise e
            print('Error: google drive Upload method failed to upload {} in {}'.format(file, folder_metadata['title']))
            return 1
    
    
    
def send_to_Googledocs(file, folder):
    """Sends document to a target folder in Google Drive account.
    
    Parameters
    ----------
    file:    Str
        Filepath for the document to be opened, read and used to create a Google Doc.
    
    folder:  Str
        Unique id for a Google Drive direcotry that allows for targeting a specific folder for creation, 
        deletion, or for adding files to a target directory. Example folder id string could look like
        the following: 169W-OsOEVpJu2vSmhhbQ2a-ZqMVCgegG as an example
    
    Returns
    -------
    int:
        Returns 0 if sucessful and 1 if process fails.
    
    
    Examples
    --------
    >>> send_to_Googledocs(file=document)
    
    """
    
    
    attempts = 0
    for attempt in range(3):
        
        try:
            drive = google_authentication()
            attempts += 1
            
            # quit trying if drive authentication succesfull 
            if drive: 
                
                break
        
        except Exception as e:
            raise e
            print(f'Google Authentication failed after {attempts} attempts.')
            
    
    try:
        googledrive_upload(drive, file, folder=folder)
        
    except Exception as e:
        raise e
        print('Error: failed to upload {} to Google Drive.'.format(file))
        return 1
    
    print('Program run sucesfully')
    return 0


###########################################################################################
                              ## Main Program Run Function ##
###########################################################################################

def run_program(division, yaml_file):
    """Executes sequence for the series of tasks and functions in the position description program. The function is a 
    wrapper for the core tasks of the program command line application for generating and handling standard formatted 
    position descriptions for the agency. 
    

    
    The function executes the following steps:
    
    1. Reads, parses a yaml file with positon data.
    2. Converts file data into dictionary object 
    3. Creates an interactive chart illustrating job responsibilities
    4. Publishes the chart to remote Plotly chart studio account, creating a unique access url
    5. Creates a directory 'data/charts' and timestamps and saves the chart as a png file to that directory
    6. Creates and saves a standard formated word document for the position description, inserting the png file to the document
    7. Logs into Google, creates or determines the existence of a folder labeled 'Position Descriptions' and pushes the word document to that Google Drive folder.
    
    Parameters
    ----------
    division:  Str
        The name of the office or division in which the position will be housed.
        
    yaml_file: Str
        Yaml file containing the information on each position. Each position in 
        the Yaml file has an attribute corresponding to the attributes for the 
        Position class. 
    
    Returns
    -------
    int:
        Returns 0 upon success and 1 upon failure.
        
    Examples
    --------
    >>> run_program(bpio,file)
    
    """
    
    folderID = None
    position_description = None
    image_file = None
    word_document = None
    send = None
    
    execution_steps = [folderID,
                       position_description,
                       image_file,
                       word_document,
                       send]
    
    ##  implement progress bar to display in terminal to track program status
    with tqdm(total = len(execution_steps), desc = 'Program Progress:', unit = 'tasks') as progress_bar:
        
        for step in execution_steps:
            
            if step == folderID and not step:
                folderID = get_googledrive_folderID(division)
                progress_bar.update(1)
            
            if step == position_description and not step:
                position_description = get_position(filename)
                progress_bar.update(1)
            
            if step == image_file and not step:
                image_file, radarchart_url = generate_plot(position_description)
                progress_bar.update(1)
                
            if step == word_document and not step:
                word_document = create_document(position_description, image_file)
                progress_bar.update(1)
            
            if step == send and not step:
                send = send_to_Googledocs(word_document, folderID)
                progress_bar.update(1)
        
    
